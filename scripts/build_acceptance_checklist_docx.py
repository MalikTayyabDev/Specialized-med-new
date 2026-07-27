"""Build Implementation Manual v2 acceptance checklist with implementation status."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT_FILE = ROOT / "docs" / "Specialized_Medical_Landing_Pages_Acceptance_Checklist.docx"

DONE = "\u2611 "  # ☑
PENDING = "\u2610 "  # ☐

PAGES = [
    {
        "num": 1,
        "name": "Cardiac Monitoring Services",
        "url": "/cardiac-monitoring-services",
        "assets": [
            ("Original service comparison graphic", True, None),
            ("Hook Up > Enroll > Monitor > Report workflow diagram", True, None),
            (
                "Portal screenshot showing report review and electronic signature",
                False,
                "Specialized Medical to provide an approved de-identified portal screenshot.",
            ),
            (
                "Original product image showing S-Patch and lead-wire options",
                True,
                "Current site photography is in place; Specialized Medical may supply updated product images.",
            ),
        ],
        "links": [
            "/mobile-cardiac-telemetry-mct",
            "/holter-monitoring-services",
            "/long-term-holter-monitoring",
            "/cardiac-event-monitoring",
            "/ambulatory-cardiac-monitoring",
            "/cardiology-practice-cardiac-monitoring",
        ],
        "schema": "Organization, MedicalBusiness, Service, WebPage, BreadcrumbList, FAQPage",
    },
    {
        "num": 2,
        "name": "Mobile Cardiac Telemetry (MCT)",
        "url": "/mobile-cardiac-telemetry-mct",
        "assets": [
            ("Wearable-to-phone-to-cloud transmission diagram", True, None),
            (
                "De-identified MCT report screenshot",
                False,
                "Specialized Medical to provide an approved de-identified MCT report image.",
            ),
            ("Phone proximity and charging illustration", True, None),
            ("Live versus non-live monitoring comparison table", True, None),
        ],
        "links": [
            "/cardiac-monitoring-services",
            "/live-ecg-monitoring",
            "/cardiac-event-monitoring",
            "/holter-monitoring-services",
            "/post-tavr-cardiac-monitoring",
            "/s-patch-cardiac-monitoring-system",
        ],
        "schema": "Service, MedicalWebPage, BreadcrumbList, FAQPage, VideoObject where video is present",
    },
    {
        "num": 3,
        "name": "Holter Monitoring Services",
        "url": "/holter-monitoring-services",
        "assets": [
            ("Patch and lead-wire Holter product images", True, None),
            ("Patient hookup sequence", True, None),
            (
                "De-identified final report sample",
                False,
                "Specialized Medical to provide an approved de-identified Holter report image.",
            ),
            ("Holter vs Extended Holter vs MCT table", True, None),
        ],
        "links": [
            "/cardiac-monitoring-services",
            "/long-term-holter-monitoring",
            "/mobile-cardiac-telemetry-mct",
            "/cardiac-event-monitoring",
            "/s-patch-cardiac-monitoring-system",
        ],
        "schema": "Service, MedicalWebPage, BreadcrumbList, FAQPage",
    },
    {
        "num": 4,
        "name": "Long-Term Holter Monitoring",
        "url": "/long-term-holter-monitoring",
        "assets": [
            ("Multi-day recording timeline", True, None),
            (
                "Full-disclosure report image",
                False,
                "Specialized Medical to provide an approved de-identified full-disclosure report image.",
            ),
            ("Electrode placement and rotation illustration", True, None),
            ("Long-Term Holter vs MCT comparison", True, None),
        ],
        "links": [
            "/holter-monitoring-services",
            "/mobile-cardiac-telemetry-mct",
            "/cardiac-event-monitoring",
            "/cardiac-monitoring-services",
            "/s-patch-cardiac-monitoring-system",
        ],
        "schema": "Service, MedicalWebPage, BreadcrumbList, FAQPage",
    },
    {
        "num": 5,
        "name": "Cardiac Event Monitoring",
        "url": "/cardiac-event-monitoring",
        "assets": [
            ("Symptom-to-event workflow diagram", True, None),
            ("Patient symptom button illustration", True, None),
            (
                "De-identified event report",
                False,
                "Specialized Medical to provide an approved de-identified event report image.",
            ),
            ("Event vs MCT vs Holter comparison", True, None),
        ],
        "links": [
            "/mobile-cardiac-telemetry-mct",
            "/holter-monitoring-services",
            "/long-term-holter-monitoring",
            "/cardiac-monitoring-services",
            "/live-ecg-monitoring",
        ],
        "schema": "Service, MedicalWebPage, BreadcrumbList, FAQPage",
    },
    {
        "num": 6,
        "name": "Ambulatory Cardiac Monitoring",
        "url": "/ambulatory-cardiac-monitoring",
        "assets": [
            ("Four-modality overview graphic", True, None),
            ("Patient/practice/monitoring-center swimlane", True, None),
            ("Signal quality and artifact illustration", True, None),
            ("Decision guide linking to each service", True, None),
        ],
        "links": [
            "/cardiac-monitoring-services",
            "/holter-monitoring-services",
            "/long-term-holter-monitoring",
            "/cardiac-event-monitoring",
            "/mobile-cardiac-telemetry-mct",
        ],
        "schema": "MedicalWebPage, ItemList, Service, BreadcrumbList, FAQPage",
    },
    {
        "num": 7,
        "name": "S-Patch Cardiac Monitoring System",
        "url": "/s-patch-cardiac-monitoring-system",
        "assets": [
            (
                "High-resolution front, side, and worn product images",
                True,
                "Current product photography is in place; Specialized Medical may supply additional high-resolution images.",
            ),
            ("Monitor-to-phone connectivity diagram", True, None),
            ("Patient setup and care infographic", True, None),
            (
                "Troubleshooting table with screenshots",
                True,
                "Troubleshooting content is implemented; Specialized Medical may supply branded screenshots if preferred.",
            ),
        ],
        "links": [
            "/mobile-cardiac-telemetry-mct",
            "/holter-monitoring-services",
            "/long-term-holter-monitoring",
            "/cardiac-event-monitoring",
            "/live-ecg-monitoring",
        ],
        "schema": "Product, Service, MedicalDevice where legally appropriate and validated, BreadcrumbList, FAQPage, VideoObject",
    },
    {
        "num": 8,
        "name": "Live ECG Monitoring",
        "url": "/live-ecg-monitoring",
        "assets": [
            ("Live data-flow animation or diagram", True, None),
            ("Connectivity/reconnection graphic", True, None),
            ("Notification workflow illustration", True, None),
            ("Live MCT vs post-study Holter comparison", True, None),
        ],
        "links": [
            "/mobile-cardiac-telemetry-mct",
            "/post-tavr-cardiac-monitoring",
            "/holter-monitoring-services",
            "/long-term-holter-monitoring",
            "/s-patch-cardiac-monitoring-system",
        ],
        "schema": "Service, MedicalWebPage, BreadcrumbList, FAQPage, VideoObject",
    },
    {
        "num": 9,
        "name": "Post-TAVR Cardiac Monitoring",
        "url": "/post-tavr-cardiac-monitoring",
        "assets": [
            ("Post-discharge workflow diagram", True, None),
            (
                "Caregiver/patient checklist",
                False,
                "Specialized Medical to confirm or supply the approved patient/caregiver checklist.",
            ),
            (
                "Sample notification protocol form",
                False,
                "Specialized Medical to confirm or supply the approved notification protocol example.",
            ),
            (
                "De-identified post-TAVR report example",
                False,
                "Specialized Medical to provide an approved de-identified post-TAVR report image.",
            ),
        ],
        "links": [
            "/mobile-cardiac-telemetry-mct",
            "/live-ecg-monitoring",
            "/cardiac-monitoring-services",
            "/cardiology-practice-cardiac-monitoring",
        ],
        "schema": "Service, MedicalWebPage, BreadcrumbList, FAQPage, HowTo only if patient steps are medically reviewed",
    },
    {
        "num": 10,
        "name": "Cardiac Monitoring for Cardiology Practices",
        "url": "/cardiology-practice-cardiac-monitoring",
        "assets": [
            ("Practice workflow swimlane", True, None),
            ("Pre-enrollment to hookup process", True, None),
            (
                "Portal and e-signature screenshots",
                False,
                "Specialized Medical to provide approved de-identified portal and e-signature screenshots.",
            ),
            ("Multi-location implementation diagram", True, None),
        ],
        "links": [
            "/cardiac-monitoring-services",
            "/ambulatory-cardiac-monitoring",
            "/mobile-cardiac-telemetry-mct",
            "/s-patch-cardiac-monitoring-system",
            "/post-tavr-cardiac-monitoring",
        ],
        "schema": "Service, MedicalBusiness, Organization, BreadcrumbList, FAQPage, VideoObject",
    },
]

GLOBAL_STANDARDS = [
    (
        "Use the supplied copy as the substantive foundation. Improve readability and page design, but do not change operational or clinical meaning without approval.",
        True,
        "Specialized Medical to confirm clinical and operational accuracy before publication.",
    ),
    (
        "Each page must answer its primary search query within the opening 100 words and then expand into a complete, authoritative resource.",
        True,
        None,
    ),
    (
        "Use one H1 only. Build a logical H2/H3 hierarchy. Do not use headings merely for visual styling.",
        True,
        None,
    ),
    (
        "Use original Specialized Medical images, diagrams, screenshots, de-identified reports, and workflow assets whenever possible. Generic stock imagery should be secondary.",
        True,
        "Specialized Medical to supply or approve final report screenshots, portal images, and any upgraded branded visuals.",
    ),
    (
        'Add descriptive internal links where specified. Avoid repetitive "click here" anchors.',
        True,
        None,
    ),
    (
        "FAQ answers must remain visible on the page even when FAQ schema is implemented.",
        True,
        None,
    ),
    (
        "All schema must match visible content. Do not mark up claims, ratings, reviews, services, products, or medical information that are not present on the page.",
        True,
        None,
    ),
    (
        "Every page must contain a clear emergency disclaimer when patient-facing content could be interpreted as urgent medical guidance.",
        True,
        None,
    ),
    (
        "Every page must pass mobile, accessibility, Core Web Vitals, metadata, indexability, canonical, and structured-data checks before launch.",
        True,
        "Specialized Medical to verify final production performance and accessibility after go-live.",
    ),
]

PAGE_ACCEPTANCE = [
    ("Primary query is answered clearly in the first 100 words.", True, None),
    (
        "All operational claims match the current Specialized Medical workflow.",
        False,
        "Specialized Medical clinical/operations review required.",
    ),
    ("Holter alert language is correct wherever compared.", True, "Specialized Medical to confirm wording."),
    (
        "Original visual assets are present and have descriptive alt text.",
        True,
        "Specialized Medical to approve final report, portal, and product images where noted above.",
    ),
    ("All required internal links are implemented and tested.", True, None),
    (
        "FAQs are visible, accurate, and marked up only when eligible.",
        True,
        "Specialized Medical to confirm FAQ clinical accuracy.",
    ),
    (
        "CTA form works on desktop and mobile and routes to the correct recipient.",
        True,
        "Specialized Medical to confirm form notifications reach the correct team.",
    ),
    (
        "Page passes spelling, HIPAA, medical, legal, schema, accessibility, and technical SEO review.",
        False,
        "Specialized Medical legal/compliance review required.",
    ),
]

LAUNCH_SECTIONS = [
    {
        "title": "Content and Accuracy",
        "items": [
            (
                "Content matches current services and device capabilities.",
                False,
                "Specialized Medical clinical and product review required.",
            ),
            (
                "No unsupported clinical claims, guarantees, or comparative superiority claims.",
                True,
                "Specialized Medical to confirm before publication.",
            ),
            (
                "All product dimensions, battery statements, connectivity statements, and service durations are verified.",
                False,
                "Specialized Medical to verify product specifications.",
            ),
            (
                "All report screenshots and ECG examples are de-identified and approved.",
                False,
                "Specialized Medical to supply and approve de-identified report images.",
            ),
        ],
    },
    {
        "title": "SEO and AEO",
        "items": [
            ("Unique title tag and meta description for every page.", True, None),
            ("One H1 per page with logical H2/H3 hierarchy.", True, None),
            (
                "Canonical URL, indexability, XML sitemap, and breadcrumb implementation verified.",
                True,
                "Specialized Medical to confirm on the live production site.",
            ),
            (
                "Internal links use descriptive anchors and point to final production URLs.",
                True,
                "Specialized Medical to confirm on the live production site.",
            ),
            ("Content provides direct answers suitable for featured snippets and AI retrieval.", True, None),
        ],
    },
    {
        "title": "Technical and Accessibility",
        "items": [
            ("Mobile layout tested at common breakpoints.", True, None),
            ("Images compressed and lazy-loaded where appropriate.", True, None),
            (
                "Core Web Vitals and page speed reviewed.",
                False,
                "Specialized Medical to review on the live production site.",
            ),
            (
                "Heading order, keyboard navigation, focus states, labels, contrast, and alt text reviewed.",
                True,
                "Specialized Medical accessibility review on production recommended.",
            ),
            ("Structured data validates and matches visible content.", True, None),
        ],
    },
    {
        "title": "Conversion and Measurement",
        "items": [
            ("Primary CTA appears above the fold and after key decision sections.", True, None),
            ("Forms have a clear success state and route to the correct team.", True, "Specialized Medical to confirm routing."),
            (
                "GA4 events are configured for form starts, form submissions, phone clicks, email clicks, video plays, and downloads.",
                False,
                "Specialized Medical to configure GA4/GTM goals in the analytics container.",
            ),
            (
                "Thank-you pages or events support conversion measurement.",
                True,
                "Specialized Medical to confirm measurement setup in GA4/GTM.",
            ),
        ],
    },
    {
        "title": "Governance",
        "items": [
            ("Page owner and review date are documented.", False, "Specialized Medical to assign page ownership."),
            ("Clinical content is reviewed periodically.", False, "Specialized Medical governance process required."),
            (
                "Changes to service capabilities are reflected across all affected pages and comparison tables.",
                False,
                "Specialized Medical to maintain content when services change.",
            ),
            (
                "No page is published with placeholder text or unverified claims.",
                True,
                "Specialized Medical final publication approval required.",
            ),
        ],
    },
]

COMPARISON_LANGUAGE = (
    "Use this comparison framework wherever monitoring modalities are compared. Every Specialized "
    "Medical test must be described as having LIVE test-status visibility while the study is in progress. "
    "This means Specialized Medical can see operational information such as battery level, electrode contact "
    "and signal quality, whether the monitor is communicating, and whether the patient appears to remain "
    "properly connected. When a parameter falls outside the expected range, Specialized Medical contacts "
    "the patient and works with the patient to correct the issue so the study has the best opportunity to "
    "be completed successfully. The distinction between test types is not whether the study is visible live; "
    "it is when clinical findings are presented. For Holter and Extended / Long-Term Holter, clinical results "
    "are presented after the final report is generated. For Event Monitoring and Mobile Cardiac Telemetry (MCT), "
    "qualifying clinical findings are presented while the test is in progress according to the prescribed "
    "notification protocol."
)

GLOBAL_CTAS = [
    ("Primary: Request a Demonstration", True, None),
    ("Secondary: Speak With a Cardiac Monitoring Specialist", True, None),
    ("Practice-focused: Review Your Current Monitoring Workflow", True, None),
    ("Post-TAVR: Request a Post-Discharge Monitoring Workflow Review", True, None),
]

CLINICAL_REVIEW = [
    (
        "All medical, technical, reimbursement, product-specification, and performance claims must be verified by Specialized Medical before publication.",
        False,
        None,
    ),
    (
        "Do not add unsupported statistics, guarantees, superiority claims, or emergency-response language.",
        True,
        "Specialized Medical to confirm compliance before publication.",
    ),
]

FINAL_DIRECTION = (
    "These pages should not be treated as thin campaign landing pages. They should function as the "
    "authoritative Specialized Medical resource library for ambulatory cardiac monitoring. The "
    "implementation team may refine wording for flow, design, and search presentation, but it should "
    "preserve the operational accuracy, modality distinctions, physician-alert timing, patient-safety "
    "language, and Specialized Medical differentiators contained in this manual."
)


def add_status_item(doc: Document, text: str, done: bool, note: str | None = None) -> None:
    prefix = DONE if done else PENDING
    paragraph = doc.add_paragraph()
    paragraph.add_run(prefix + text)
    if note:
        run = paragraph.add_run(f"  ({note})")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_status_items(doc: Document, items: list[tuple[str, bool, str | None]]) -> None:
    for text, done, note in items:
        add_status_item(doc, text, done, note)


def main() -> None:
    OUT_FILE.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("SPECIALIZED MEDICAL", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_heading("SEO Landing Page Content & Implementation Manual", level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Acceptance Checklist").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Prepared for Specialized Medical LLC").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("July 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_paragraph(
        "This checklist follows the implementation manual for all ten landing pages. "
        "Items marked complete reflect work already implemented. "
        "Items marked pending require Specialized Medical review, approval, asset delivery, or production verification."
    )
    doc.add_paragraph()
    doc.add_paragraph(f"{DONE}Complete")
    doc.add_paragraph(f"{PENDING}Pending — Specialized Medical action required")
    doc.add_paragraph()

    doc.add_heading("Clinical and Compliance Review", level=1)
    add_status_items(doc, CLINICAL_REVIEW)

    doc.add_heading("1. Implementation Standards for All Pages", level=1)
    add_status_items(doc, GLOBAL_STANDARDS)

    doc.add_heading("Required Global Comparison Language", level=2)
    doc.add_paragraph(COMPARISON_LANGUAGE)
    add_status_item(
        doc,
        "Comparison framework and modality table implemented on all relevant pages.",
        True,
        "Specialized Medical to confirm clinical wording.",
    )

    doc.add_heading("Global Calls to Action", level=2)
    add_status_items(doc, GLOBAL_CTAS)

    doc.add_heading("Global Disclaimers", level=2)
    doc.add_paragraph("Diagnostic Service Disclaimer")
    doc.add_paragraph(
        "Specialized Medical provides ambulatory cardiac monitoring services and diagnostic information "
        "for review by qualified healthcare providers. The service does not replace physician judgment, "
        "emergency medical services, or instructions to call 911 when urgent symptoms occur."
    )
    add_status_item(doc, "Diagnostic Service Disclaimer present on all applicable pages.", True, None)
    add_status_item(
        doc,
        "Emergency guidance language present on all patient-facing pages.",
        True,
        "Specialized Medical to confirm wording.",
    )

    doc.add_page_break()
    doc.add_heading("2. Page-by-Page Content Package", level=1)

    for page in PAGES:
        doc.add_heading(f"{page['num']}. {page['name']}", level=2)
        doc.add_paragraph(f"URL {page['url']}")
        add_status_item(doc, "Landing page built with manual copy, metadata, FAQs, CTA, and schema.", True, None)

        doc.add_heading("Required Original Assets", level=3)
        add_status_items(doc, page["assets"])

        doc.add_heading("Required Internal Links", level=3)
        for link in page["links"]:
            add_status_item(doc, link, True, None)

        doc.add_heading("Recommended Structured Data", level=3)
        add_status_item(doc, page["schema"], True, None)

        doc.add_heading("Page Acceptance Checklist", level=3)
        add_status_items(doc, PAGE_ACCEPTANCE)
        doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading("13. Launch and Acceptance Checklist", level=1)

    for section in LAUNCH_SECTIONS:
        doc.add_heading(section["title"], level=2)
        add_status_items(doc, section["items"])

    doc.add_heading("Final Direction to the SEO Team", level=2)
    doc.add_paragraph(FINAL_DIRECTION)

    doc.add_paragraph()
    doc.add_paragraph("All ten landing pages reviewed and accepted by Specialized Medical:")
    doc.add_paragraph()
    doc.add_paragraph("Name: _________________________________    Title: _________________________________")
    doc.add_paragraph("Signature: _____________________________    Date: _________________________________")

    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    main()

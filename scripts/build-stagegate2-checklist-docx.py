#!/usr/bin/env python3
"""Build Stagegate 2 acceptance checklist Word document from markdown source."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "08-stagegate-2-acceptance-checklist.md"
OUTPUT = ROOT / "docs" / "Specialized_Medical_Stagegate_2_Checklist.docx"


def strip_md_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1 (\2)", text)
    return text.strip()


def add_rich_paragraph(doc: Document, text: str, style: str | None = None):
    para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(strip_md_inline(part))
    return para


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [strip_md_inline(cell.strip()) for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line.strip()))


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        from docx.oxml import OxmlElement

        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def style_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    for level in (1, 2, 3):
        heading = doc.styles[f"Heading {level}"]
        heading.font.name = "Calibri"
        heading.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return

    col_count = max(len(row) for row in rows)
    normalized = [row + [""] * (col_count - len(row)) for row in rows]

    table = doc.add_table(rows=len(normalized), cols=col_count)
    table.style = "Table Grid"
    table.autofit = True

    for r_idx, row in enumerate(normalized):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            run.font.size = Pt(9.5 if r_idx else 10)
            if r_idx == 0:
                run.bold = True
                set_cell_shading(cell, "D9E2F3")
            if c_idx == 0 and r_idx > 0 and cell_text in ("[x]", "[ ]"):
                run.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x70, 0x30) if cell_text == "[x]" else RGBColor(0xC0, 0x00, 0x00)


def build_docx(source: Path, output: Path):
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style_document(doc)

    table_buffer: list[list[str]] = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            if table_buffer:
                add_table(doc, table_buffer)
                table_buffer = []
            i += 1
            continue

        if stripped == "---":
            if table_buffer:
                add_table(doc, table_buffer)
                table_buffer = []
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("|"):
            if is_table_separator(stripped):
                i += 1
                continue
            table_buffer.append(parse_table_row(stripped))
            i += 1
            continue

        if table_buffer:
            add_table(doc, table_buffer)
            table_buffer = []

        if stripped.startswith("# "):
            title = doc.add_heading(strip_md_inline(stripped[2:]), level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith("## "):
            doc.add_heading(strip_md_inline(stripped[3:]), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(strip_md_inline(stripped[4:]), level=2)
        elif stripped.startswith("- "):
            add_rich_paragraph(doc, stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            add_rich_paragraph(doc, re.sub(r"^\d+\.\s", "", stripped), style="List Number")
        else:
            add_rich_paragraph(doc, stripped)

        i += 1

    if table_buffer:
        add_table(doc, table_buffer)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Wrote {output}")


if __name__ == "__main__":
    src = Path(sys.argv[1]) if len(sys.argv) > 1 else SOURCE
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else OUTPUT
    build_docx(src, out)

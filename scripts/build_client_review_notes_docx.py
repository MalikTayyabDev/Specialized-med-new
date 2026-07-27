"""Client-facing notes on PDF manual items requiring Specialized Medical input."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
OUT_FILE = ROOT / "docs" / "Specialized_Medical_Client_Review_Notes.docx"

SECTIONS = [
    {
        "title": "Purpose",
        "paragraphs": [
            "The ten landing pages have been rebuilt to follow the SEO Content & Implementation Manual (July 2026) as closely as possible. "
            "Publication-ready copy, metadata, FAQs, comparison language, CTAs, disclaimers, internal links, and schema follow the manual.",
            "The items below are points where the manual is incomplete, internally inconsistent, or depends on Specialized Medical assets or approval. "
            "Please review and confirm so the pages can be signed off.",
        ],
    },
    {
        "title": "1. Meta descriptions truncated in the manual",
        "paragraphs": [
            "Several recommended meta description fields in the manual end with an ellipsis (...). "
            "Where that occurs, the live pages complete the sentence using only the next words from the Publication-Ready Opening Copy on the same page. "
            "No new marketing language was added.",
        ],
        "items": [
            "Long-Term Holter Monitoring: the meta field ends with “Specialized Medical can support...” while the opening copy continues with “Specialized Medical maintains LIVE test-status visibility...” Please confirm the preferred meta description.",
            "Please confirm final meta descriptions for all ten pages before production publication.",
        ],
    },
    {
        "title": "2. Internal inconsistency in the manual (LIVE language)",
        "paragraphs": [
            "Section 1 and the global comparison framework state that every test type has LIVE test-status visibility during the study.",
            "The Long-Term Holter FAQ in the manual asks “Does Long-Term Holter provide live alerts?” and answers “No. The study is reviewed after completion, and the physician is alerted after final reports are generated.”",
            "The Holter Monitoring page FAQ uses the LIVE operational visibility wording consistent with Section 1.",
            "The live pages currently use the manual FAQ text exactly as written on each page. Please confirm which wording is clinically correct so all pages stay consistent.",
        ],
    },
    {
        "title": "3. Comparison table wording",
        "paragraphs": [
            "The manual uses two phrasings for Holter clinical-result timing:",
        ],
        "items": [
            "Table cells: “After Final Report Is Generated”",
            "Body instructions on some pages: “After Final Reports Generated”",
            "Live pages use the table wording. Please confirm if one phrasing should be used everywhere.",
        ],
    },
    {
        "title": "4. Original assets not supplied in the manual",
        "paragraphs": [
            "The manual lists required original images, portal screenshots, and de-identified reports. "
            "Until Specialized Medical supplies approved files, the pages use existing site photography, a de-identified report placeholder, and simple HTML workflow illustrations that match the manual’s asset descriptions.",
        ],
        "items": [
            "Approved de-identified Holter, MCT, Event, and Post-TAVR report screenshots",
            "Approved portal and electronic-signature screenshots",
            "High-resolution S-Patch product photography (front, side, worn) if different from current site images",
            "Branded workflow diagrams if HTML placeholders should be replaced",
            "Post-TAVR caregiver/patient checklist and sample notification protocol form (if these should appear as downloadable or embedded assets)",
            "Any video files referenced for VideoObject schema",
        ],
    },
    {
        "title": "5. Product specifications (S-Patch page)",
        "paragraphs": [
            "The manual lists approximate disk dimensions and states that battery life, water resistance, and connectivity claims must be verified before publication. "
            "Please confirm all product specifications currently on the S-Patch page.",
        ],
    },
    {
        "title": "6. Clinical, legal, and operational sign-off",
        "items": [
            "Verify all medical, workflow, reimbursement, and performance statements match current Specialized Medical services",
            "Confirm FAQ answers on every page",
            "Confirm emergency and diagnostic disclaimers",
            "Confirm CTA form notifications reach the correct team",
            "Assign page owner and review date for governance (Section 13)",
        ],
    },
    {
        "title": "7. Analytics and production verification",
        "items": [
            "Configure GA4/GTM goals for form starts, submissions, phone clicks, email clicks, video plays, and downloads (Section 13)",
            "Review all ten pages on the live production domain after deploy",
            "Confirm Core Web Vitals and accessibility on production",
        ],
    },
    {
        "title": "8. What was implemented from the manual (no client action needed unless noted above)",
        "items": [
            "All ten landing pages with Publication-Ready opening copy, body sections, and headings from the manual",
            "Ten FAQs per page with manual wording",
            "Primary CTA labels and form fields specified in each CTA block",
            "Global comparison language and modality table",
            "Required internal links and recommended structured data",
            "Diagnostic Service Disclaimer and patient-facing emergency language where required",
            "XML sitemap, breadcrumbs, canonical URLs, and indexability",
        ],
    },
]


def main() -> None:
    OUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Specialized Medical", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_heading("Landing Page Implementation — Client Review Notes", level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Prepared for Specialized Medical LLC").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("July 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for section in SECTIONS:
        doc.add_heading(section["title"], level=1)
        for paragraph in section.get("paragraphs", []):
            doc.add_paragraph(paragraph)
        for item in section.get("items", []):
            doc.add_paragraph(item, style="List Bullet")
        doc.add_paragraph()

    doc.add_paragraph("Specialized Medical reviewer: _________________________________    Date: _________________")

    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    main()
